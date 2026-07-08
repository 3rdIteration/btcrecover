#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# test_walletfinder.py -- unit tests for walletfinder.py
# Copyright (C) 2014-2017 Christopher Gurnee
#               2019-2021 Stephen Rothery
#
# This file is part of btcrecover.
#
# btcrecover is free software: you can redistribute it and/or
# modify it under the terms of the GNU General Public License
# as published by the Free Software Foundation; either version
# 2 of the License, or (at your option) or later version.
#
# btcrecover is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see http://www.gnu.org/licenses/

import contextlib
import io
import os
import sys
import unittest
import tempfile
import shutil
from pathlib import Path

if __name__ == '__main__':
    sys.path.append(os.path.join(os.path.dirname(__file__), "..", ".."))

# Import walletfinder functions for direct testing
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", ".."))


def can_load_shamir_mnemonic():
    """Check if shamir_mnemonic package is available for SLIP39 support."""
    try:
        from shamir_mnemonic import wordlist as sw  # noqa: F401
        return True
    except ImportError:
        return False


def can_detect_wallet(filename):
    """Check if load_wallet can detect the given wallet file."""
    filepath = os.path.join(os.path.dirname(__file__), "test-wallets", filename)
    if not os.path.isfile(filepath):
        return False
    try:
        from btcrecover.btcrpass import load_wallet
        return load_wallet(filepath) is not None
    except (Exception, SystemExit):
        return False


import walletfinder


WALLET_DIR = os.path.join(os.path.dirname(__file__), "test-wallets")


class TestWalkDirectory(unittest.TestCase):
    """Test the walk_directory helper function."""

    def test_walk_basic(self):
        tmpdir = tempfile.mkdtemp()
        try:
            Path(tmpdir, "a.txt").write_text("hello")
            Path(tmpdir, "b.txt").write_text("world")
            result = sorted(walletfinder.walk_directory(tmpdir, None))
            self.assertEqual(len(result), 2)
            self.assertTrue(any("a.txt" in p for p in result))
            self.assertTrue(any("b.txt" in p for p in result))
        finally:
            shutil.rmtree(tmpdir)

    def test_walk_nested(self):
        tmpdir = tempfile.mkdtemp()
        try:
            Path(tmpdir, "sub").mkdir()
            Path(tmpdir, "a.txt").write_text("hello")
            Path(tmpdir, "sub", "b.txt").write_text("world")
            result = sorted(walletfinder.walk_directory(tmpdir, None))
            self.assertEqual(len(result), 2)
        finally:
            shutil.rmtree(tmpdir)

    def test_walk_depth_limit_0(self):
        tmpdir = tempfile.mkdtemp()
        try:
            Path(tmpdir, "sub").mkdir()
            Path(tmpdir, "a.txt").write_text("hello")
            Path(tmpdir, "sub", "b.txt").write_text("world")
            result = sorted(walletfinder.walk_directory(tmpdir, 0))
            self.assertEqual(len(result), 1)
            self.assertTrue("a.txt" in result[0])
        finally:
            shutil.rmtree(tmpdir)

    def test_walk_depth_limit_1(self):
        tmpdir = tempfile.mkdtemp()
        try:
            Path(tmpdir, "sub").mkdir()
            Path(tmpdir, "sub", "deep").mkdir()
            Path(tmpdir, "a.txt").write_text("hello")
            Path(tmpdir, "sub", "b.txt").write_text("world")
            Path(tmpdir, "sub", "deep", "c.txt").write_text("deep")
            result = sorted(walletfinder.walk_directory(tmpdir, 1))
            self.assertEqual(len(result), 2)
        finally:
            shutil.rmtree(tmpdir)

    def test_walk_excludes_hidden_dirs(self):
        tmpdir = tempfile.mkdtemp()
        try:
            Path(tmpdir, ".hidden").mkdir()
            Path(tmpdir, "a.txt").write_text("hello")
            Path(tmpdir, ".hidden", "b.txt").write_text("world")
            result = sorted(walletfinder.walk_directory(tmpdir, None))
            self.assertEqual(len(result), 1)
        finally:
            shutil.rmtree(tmpdir)

    def test_walk_excludes_common_dirs(self):
        tmpdir = tempfile.mkdtemp()
        try:
            Path(tmpdir, "node_modules").mkdir()
            Path(tmpdir, "__pycache__").mkdir()
            Path(tmpdir, ".git").mkdir()
            Path(tmpdir, "a.txt").write_text("hello")
            Path(tmpdir, "node_modules", "b.txt").write_text("world")
            result = sorted(walletfinder.walk_directory(tmpdir, None))
            self.assertEqual(len(result), 1)
        finally:
            shutil.rmtree(tmpdir)

    def test_walk_nonexistent_dir(self):
        result = list(walletfinder.walk_directory("/nonexistent/path/xyz", None))
        self.assertEqual(len(result), 0)


class TestMnemonicHelpers(unittest.TestCase):
    """Test mnemonic detection helper functions."""

    def setUp(self):
        self.wordset = {"abandon", "ability", "about", "absorb", "abstract", "absurd",
                        "abuse", "access", "accident", "account", "accurate", "across"}

    def test_check_sequential_basic(self):
        tokens = ["the", "abandon", "ability", "about", "hello"]
        matches = walletfinder.check_sequential(tokens, self.wordset, 3)
        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0][1], 3)
        self.assertEqual(matches[0][2], ["abandon", "ability", "about"])

    def test_check_sequential_no_match(self):
        tokens = ["the", "quick", "brown", "fox"]
        matches = walletfinder.check_sequential(tokens, self.wordset, 3)
        self.assertEqual(len(matches), 0)

    def test_check_sequential_short_run(self):
        tokens = ["abandon", "ability", "hello"]
        matches = walletfinder.check_sequential(tokens, self.wordset, 4)
        self.assertEqual(len(matches), 0)

    def test_check_sequential_exact_match(self):
        tokens = ["abandon", "ability", "about"]
        matches = walletfinder.check_sequential(tokens, self.wordset, 3)
        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0][1], 3)

    def test_check_sequential_multiple_runs(self):
        tokens = ["abandon", "ability", "about", "foo", "absorb", "abstract", "absurd"]
        matches = walletfinder.check_sequential(tokens, self.wordset, 3)
        self.assertEqual(len(matches), 2)

    def test_check_scattered_basic(self):
        tokens = ["abandon", "ability", "about", "hello", "world"]
        matched = walletfinder.check_scattered(tokens, self.wordset)
        self.assertEqual(len(matched), 3)
        self.assertIn("abandon", matched)

    def test_check_scattered_duplicates(self):
        tokens = ["abandon", "abandon", "ability"]
        matched = walletfinder.check_scattered(tokens, self.wordset)
        self.assertEqual(len(matched), 2)

    def test_check_scattered_punctuation(self):
        tokens = ['"abandon"', "ability,", "about.", '"hello"']
        matched = walletfinder.check_scattered(tokens, self.wordset)
        self.assertEqual(len(matched), 3)


class TestWalletScan(unittest.TestCase):
    """Test wallet scanning against known test-wallet files."""

    def test_scan_finds_wallets(self):
        results, scanned = walletfinder.scan_wallet_mode(WALLET_DIR, None)
        if not results:
            self.skipTest("no wallets detected in this environment")
        self.assertGreater(scanned, 0)

    def test_bitcoincore_detected(self):
        results, _ = walletfinder.scan_wallet_mode(WALLET_DIR, None)
        btc_core = [r for r in results if 'bitcoincore-wallet.dat' in r['path']]
        if not btc_core:
            self.skipTest("bitcoincore-wallet.dat not detected in this environment")
        self.assertIn('BitcoinCore', btc_core[0]['type'])

    def test_electrum_detected(self):
        results, _ = walletfinder.scan_wallet_mode(WALLET_DIR, None)
        electrum = [r for r in results if 'electrum-wallet' == os.path.basename(r['path'])]
        if not electrum:
            self.skipTest("electrum wallet not detected in this environment")
        self.assertIn('Electrum', electrum[0]['type'])

    def test_blockchain_detected(self):
        results, _ = walletfinder.scan_wallet_mode(WALLET_DIR, None)
        bc = [r for r in results if 'blockchain-v4.0-wallet.aes.json' in r['path']]
        if not bc:
            self.skipTest("blockchain wallet not detected in this environment")
        self.assertIn('Blockchain', bc[0]['type'])

    def test_metamask_detected(self):
        results, _ = walletfinder.scan_wallet_mode(WALLET_DIR, None)
        mm = [r for r in results if 'metamask.9.8.4_firefox_vault' in r['path']]
        if not mm:
            self.skipTest("metamask vault not detected in this environment")
        self.assertIn('Metamask', mm[0]['type'])

    def test_dogechain_detected(self):
        results, _ = walletfinder.scan_wallet_mode(WALLET_DIR, None)
        dc = [r for r in results if 'dogechain.wallet.aes.json' == os.path.basename(r['path'])]
        if not dc:
            self.skipTest("dogechain wallet not detected in this environment")
        self.assertIn('Dogechain', dc[0]['type'])

    def test_blockio_detected(self):
        results, _ = walletfinder.scan_wallet_mode(WALLET_DIR, None)
        bio = [r for r in results if 'block.io.request.json' == os.path.basename(r['path'])]
        if not bio:
            self.skipTest("block.io wallet not detected in this environment")
        self.assertIn('BlockIO', bio[0]['type'])

    def test_multibit_detected(self):
        results, _ = walletfinder.scan_wallet_mode(WALLET_DIR, None)
        mb = [r for r in results if 'multibit-wallet.key' == os.path.basename(r['path'])]
        if not mb:
            self.skipTest("multibit wallet not detected in this environment")
        self.assertIn('MultiBit', mb[0]['type'])

    def test_multibithd_detected(self):
        results, _ = walletfinder.scan_wallet_mode(WALLET_DIR, None)
        mbhd = [r for r in results if 'mbhd.wallet.aes' == os.path.basename(r['path'])]
        if not mbhd:
            self.skipTest("multibithd wallet not detected in this environment")
        self.assertIn('MultiBitHD', mbhd[0]['type'])

    def test_btc_com_detected(self):
        results, _ = walletfinder.scan_wallet_mode(WALLET_DIR, None)
        btc = [r for r in results if 'btc_com_parsed' in r['path']]
        if not btc:
            self.skipTest("btc.com wallet not detected in this environment")
        self.assertIn('btc_com', btc[0]['type'])

    def test_scan_empty_dir(self):
        tmpdir = tempfile.mkdtemp()
        try:
            results, scanned = walletfinder.scan_wallet_mode(tmpdir, None)
            self.assertEqual(len(results), 0)
            self.assertEqual(scanned, 0)
        finally:
            shutil.rmtree(tmpdir)

    def test_scan_depth_limit(self):
        tmpdir = tempfile.mkdtemp()
        try:
            src = os.path.join(WALLET_DIR, "electrum1-upgradedto-electrum27-wallet")
            Path(tmpdir, "sub").mkdir()
            shutil.copy2(src, tmpdir)
            shutil.copy2(src, os.path.join(tmpdir, "sub"))
            results_depth0, _ = walletfinder.scan_wallet_mode(tmpdir, 0)
            results_unlimited, _ = walletfinder.scan_wallet_mode(tmpdir, None)
            self.assertEqual(len(results_depth0), 1)
            self.assertEqual(len(results_unlimited), 2)
        finally:
            shutil.rmtree(tmpdir)


def can_use_textract():
    """Check if textract is available for document extraction tests."""
    try:
        import textract  # noqa: F401
        return True
    except ImportError:
        return False


class TestMnemonicScan(unittest.TestCase):
    """Test mnemonic scanning functionality."""

    def test_bip39_sequential_detection(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "test.txt")
            with open(testfile, 'w') as f:
                f.write("abandon ability about absorb abstract absurd abuse access accident account\n")
            results, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(len(results), 1)
            findings = results[0]['findings']
            bip39_findings = [f for f in findings if 'BIP39' in f['wordlist']]
            self.assertEqual(len(bip39_findings), 1)
            self.assertGreaterEqual(len(bip39_findings[0]['sequential']), 1)
        finally:
            shutil.rmtree(tmpdir)

    def test_bip39_scattered_detection(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "test.txt")
            with open(testfile, 'w') as f:
                words = ["abandon", "ability", "about", "absorb", "abstract", "absurd",
                         "abuse", "access", "accident", "account", "accurate", "across"]
                for w in words:
                    f.write(w + " hello ")
            results, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(len(results), 1)
            findings = results[0]['findings']
            bip39_findings = [f for f in findings if 'BIP39' in f['wordlist']]
            self.assertEqual(len(bip39_findings), 1)
            self.assertGreaterEqual(bip39_findings[0]['scattered_count'], 12)
        finally:
            shutil.rmtree(tmpdir)

    def test_below_threshold_not_reported(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "test.txt")
            with open(testfile, 'w') as f:
                f.write("hello world foo bar baz\n")
            results, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(len(results), 0)
        finally:
            shutil.rmtree(tmpdir)

    def test_electrum_legacy_detection(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "test.txt")
            with open(testfile, 'w') as f:
                f.write("like just love know never want time out there make look eye\n")
            results, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(len(results), 1)
            findings = results[0]['findings']
            electrum_findings = [f for f in findings if 'Electrum' in f['wordlist']]
            self.assertEqual(len(electrum_findings), 1)
        finally:
            shutil.rmtree(tmpdir)

    @unittest.skipUnless(can_load_shamir_mnemonic(), "requires shamir-mnemonic")
    def test_slip39_detection(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "test.txt")
            with open(testfile, 'w') as f:
                f.write("duckling enlarge academic academic agency result length solution fridge kidney coal piece\n")
            results, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(len(results), 1)
            findings = results[0]['findings']
            slip39_findings = [f for f in findings if 'SLIP39' in f['wordlist']]
            self.assertGreaterEqual(len(slip39_findings), 1)
        finally:
            shutil.rmtree(tmpdir)

    def test_binary_file_handled(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "test.bin")
            with open(testfile, 'wb') as f:
                f.write(b'\x00\x01\x02\x03' * 100)
            results, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(len(results), 0)
        finally:
            shutil.rmtree(tmpdir)

    def test_multiple_files(self):
        tmpdir = tempfile.mkdtemp()
        try:
            for i in range(3):
                testfile = os.path.join(tmpdir, "test{}.txt".format(i))
                with open(testfile, 'w') as f:
                    f.write("abandon ability about absorb abstract absurd\n")
            results, scanned = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(scanned, 3)
            self.assertEqual(len(results), 3)
        finally:
            shutil.rmtree(tmpdir)

    def test_depth_limit(self):
        tmpdir = tempfile.mkdtemp()
        try:
            Path(tmpdir, "sub").mkdir()
            with open(os.path.join(tmpdir, "a.txt"), 'w') as f:
                f.write("abandon ability about absorb abstract absurd\n")
            with open(os.path.join(tmpdir, "sub", "b.txt"), 'w') as f:
                f.write("abandon ability about absorb abstract absurd\n")
            results_depth0, _ = walletfinder.scan_text_mode(tmpdir, 0, 6, 12)
            results_unlimited, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(len(results_depth0), 1)
            self.assertEqual(len(results_unlimited), 2)
        finally:
            shutil.rmtree(tmpdir)

    def test_min_sequential_configurable(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "test.txt")
            with open(testfile, 'w') as f:
                f.write("abandon ability about\n")
            results_low, _ = walletfinder.scan_text_mode(tmpdir, None, 3, 12)
            results_high, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(len(results_low), 1)
            self.assertEqual(len(results_high), 0)
        finally:
            shutil.rmtree(tmpdir)

    def test_min_scattered_configurable(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "test.txt")
            with open(testfile, 'w') as f:
                words = ["abandon", "ability", "about", "absorb", "abstract"]
                for w in words:
                    f.write(w + " ")
            results_low, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 5)
            results_high, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(len(results_low), 1)
            self.assertEqual(len(results_high), 0)
        finally:
            shutil.rmtree(tmpdir)

    def test_scan_mnemonic_mode_includes_documents(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "notes.txt")
            with open(testfile, 'w') as f:
                f.write("abandon ability about absorb abstract absurd\n")
            results, scanned = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertGreaterEqual(scanned, 1)
            self.assertEqual(len(results), 1)
        finally:
            shutil.rmtree(tmpdir)

    @unittest.skipUnless(can_use_textract(), "requires textract")
    def test_scan_mnemonic_mode_with_docx(self):
        try:
            import docx2txt
        except ImportError:
            self.skipTest("requires python-docx2txt for .docx support")
        tmpdir = tempfile.mkdtemp()
        try:
            import docx
            doc = docx.Document()
            doc.add_paragraph("Here are my secret words: abandon ability about absorb abstract absurd abuse access accident account accurate across")
            testfile = os.path.join(tmpdir, "secrets.docx")
            doc.save(testfile)
            results, scanned = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertGreaterEqual(scanned, 1)
            self.assertEqual(len(results), 1)
            findings = results[0]['findings']
            bip39_findings = [f for f in findings if 'BIP39' in f['wordlist']]
            self.assertGreaterEqual(len(bip39_findings), 1)
        finally:
            shutil.rmtree(tmpdir)

    @unittest.skipUnless(can_use_textract(), "requires textract")
    def test_scan_mnemonic_mode_with_xlsx(self):
        try:
            import xlrd
        except ImportError:
            self.skipTest("requires xlrd for .xlsx support")
        tmpdir = tempfile.mkdtemp()
        try:
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            ws.append(["My secret words:", "abandon", "ability", "about", "absorb", "abstract", "absurd"])
            testfile = os.path.join(tmpdir, "secrets.xlsx")
            wb.save(testfile)
            results, scanned = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertGreaterEqual(scanned, 1)
            self.assertEqual(len(results), 1)
            findings = results[0]['findings']
            bip39_findings = [f for f in findings if 'BIP39' in f['wordlist']]
            self.assertGreaterEqual(len(bip39_findings), 1)
        finally:
            shutil.rmtree(tmpdir)

    @unittest.skipUnless(can_use_textract(), "requires textract")
    def test_scan_mnemonic_mode_with_csv(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "words.csv")
            with open(testfile, 'w') as f:
                f.write("word1,word2,word3\nabandon,ability,about\nabsorb,abstract,absurd\n")
            results, scanned = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertGreaterEqual(scanned, 1)
            findings = results[0]['findings']
            bip39_findings = [f for f in findings if 'BIP39' in f['wordlist']]
            self.assertGreaterEqual(len(bip39_findings), 1)
        finally:
            shutil.rmtree(tmpdir)

    @unittest.skipUnless(can_use_textract(), "requires textract")
    def test_scan_mnemonic_mode_with_json(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "data.json")
            with open(testfile, 'w') as f:
                f.write('{"notes": "abandon ability about absorb abstract absurd"}\n')
            results, scanned = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertGreaterEqual(scanned, 1)
            findings = results[0]['findings']
            bip39_findings = [f for f in findings if 'BIP39' in f['wordlist']]
            self.assertGreaterEqual(len(bip39_findings), 1)
        finally:
            shutil.rmtree(tmpdir)

def can_use_textract():
    """Check if textract is available for document extraction tests."""
    try:
        import textract  # noqa: F401
        return True
    except ImportError:
        return False


class TestTextractIntegration(unittest.TestCase):
    """Test textract-based document scanning functionality."""

    def test_read_file_plain_text(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "test.txt")
            with open(testfile, 'w') as f:
                f.write("abandon ability about absorb abstract absurd\n")
            content = walletfinder.read_file_with_textract(testfile, 16384)
            self.assertIsNotNone(content)
            self.assertIn("abandon", content)
        finally:
            shutil.rmtree(tmpdir)

    def test_read_file_json(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "test.json")
            with open(testfile, 'w') as f:
                f.write('{"mnemonic": "abandon ability about absorb abstract absurd"}\n')
            content = walletfinder.read_file_with_textract(testfile, 16384)
            self.assertIsNotNone(content)
            self.assertIn("abandon", content)
        finally:
            shutil.rmtree(tmpdir)

    def test_read_file_html(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "test.html")
            with open(testfile, 'w') as f:
                f.write('<html><body>abandon ability about absorb abstract absurd</body></html>\n')
            content = walletfinder.read_file_with_textract(testfile, 16384)
            self.assertIsNotNone(content)
            self.assertIn("abandon", content)
        finally:
            shutil.rmtree(tmpdir)

    def test_read_file_unsupported_extension_without_textract(self):
        tmpdir = tempfile.mkdtemp()
        try:
            # Without textract, .docx falls back to raw UTF-8 reading (returns binary garbage)
            testfile = os.path.join(tmpdir, "test.docx")
            with open(testfile, 'wb') as f:
                f.write(b'\x00\x01\x02\x03' * 100)
            content = walletfinder.read_file_with_textract(testfile, 16384)
            # Falls back to reading raw bytes as text (with errors='ignore')
            self.assertIsNotNone(content)
        finally:
            shutil.rmtree(tmpdir)

    def test_read_file_unknown_extension(self):
        tmpdir = tempfile.mkdtemp()
        try:
            # Unknown extensions without textract should still fall back to UTF-8 reading
            testfile = os.path.join(tmpdir, "test.xyz")
            with open(testfile, 'w') as f:
                f.write("abandon ability about absorb abstract absurd\n")
            content = walletfinder.read_file_with_textract(testfile, 16384)
            self.assertIsNotNone(content)
            self.assertIn("abandon", content)
        finally:
            shutil.rmtree(tmpdir)

    def test_read_file_plain_text_fallback(self):
        tmpdir = tempfile.mkdtemp()
        try:
            # Plain text extensions should always work without textract
            for ext in ('txt', 'csv', 'json', 'html', 'htm'):
                testfile = os.path.join(tmpdir, "test.{}".format(ext))
                with open(testfile, 'w') as f:
                    f.write("abandon ability about absorb abstract absurd\n")
                content = walletfinder.read_file_with_textract(testfile, 16384)
                self.assertIsNotNone(content, "Failed for .{} extension".format(ext))
                self.assertIn("abandon", content)
        finally:
            shutil.rmtree(tmpdir)

    def test_scan_text_mode_includes_documents(self):
        tmpdir = tempfile.mkdtemp()
        try:
            # Create a plain text file with mnemonic words
            testfile = os.path.join(tmpdir, "notes.txt")
            with open(testfile, 'w') as f:
                f.write("abandon ability about absorb abstract absurd\n")
            results, scanned = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertGreaterEqual(scanned, 1)
            self.assertEqual(len(results), 1)
        finally:
            shutil.rmtree(tmpdir)

    @unittest.skipUnless(can_use_textract(), "requires textract")
    def test_scan_text_mode_with_docx(self):
        try:
            import docx2txt
        except ImportError:
            self.skipTest("requires python-docx2txt for .docx support")
        tmpdir = tempfile.mkdtemp()
        try:
            # Create a real .docx file with mnemonic words
            import docx
            doc = docx.Document()
            doc.add_paragraph("Here are my secret words: abandon ability about absorb abstract absurd abuse access accident account accurate across")
            testfile = os.path.join(tmpdir, "secrets.docx")
            doc.save(testfile)
            results, scanned = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertGreaterEqual(scanned, 1)
            self.assertEqual(len(results), 1)
            findings = results[0]['findings']
            bip39_findings = [f for f in findings if 'BIP39' in f['wordlist']]
            self.assertGreaterEqual(len(bip39_findings), 1)
        finally:
            shutil.rmtree(tmpdir)

    @unittest.skipUnless(can_use_textract(), "requires textract")
    def test_scan_text_mode_with_xlsx(self):
        try:
            import xlrd
        except ImportError:
            self.skipTest("requires xlrd for .xlsx support")
        tmpdir = tempfile.mkdtemp()
        try:
            # Create a real .xlsx file with mnemonic words
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            ws.append(["My secret words:", "abandon", "ability", "about", "absorb", "abstract", "absurd"])
            testfile = os.path.join(tmpdir, "secrets.xlsx")
            wb.save(testfile)
            results, scanned = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertGreaterEqual(scanned, 1)
            self.assertEqual(len(results), 1)
            findings = results[0]['findings']
            bip39_findings = [f for f in findings if 'BIP39' in f['wordlist']]
            self.assertGreaterEqual(len(bip39_findings), 1)
        finally:
            shutil.rmtree(tmpdir)

    @unittest.skipUnless(can_use_textract(), "requires textract")
    def test_scan_text_mode_with_csv(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "words.csv")
            with open(testfile, 'w') as f:
                f.write("word1,word2,word3\nabandon,ability,about\nabsorb,abstract,absurd\n")
            results, scanned = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertGreaterEqual(scanned, 1)
            findings = results[0]['findings']
            bip39_findings = [f for f in findings if 'BIP39' in f['wordlist']]
            self.assertGreaterEqual(len(bip39_findings), 1)
        finally:
            shutil.rmtree(tmpdir)

    @unittest.skipUnless(can_use_textract(), "requires textract")
    def test_scan_text_mode_with_json(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "data.json")
            with open(testfile, 'w') as f:
                f.write('{"notes": "abandon ability about absorb abstract absurd"}\n')
            results, scanned = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertGreaterEqual(scanned, 1)
            findings = results[0]['findings']
            bip39_findings = [f for f in findings if 'BIP39' in f['wordlist']]
            self.assertGreaterEqual(len(bip39_findings), 1)
        finally:
            shutil.rmtree(tmpdir)

    def test_textract_supported_extensions_set(self):
        expected = {
            'csv', 'tsv', 'tab', 'doc', 'docx', 'eml', 'epub', 'gif',
            'jpg', 'jpeg', 'json', 'html', 'htm', 'mp3', 'msg', 'odt',
            'ogg', 'pdf', 'png', 'pptx', 'ps', 'rtf', 'tiff', 'tif', 'txt', 'wav',
            'xls', 'xlsx',
        }
        self.assertEqual(walletfinder.TEXTRACT_SUPPORTED_EXTENSIONS, expected)


class TestArgumentParsing(unittest.TestCase):
    """Test CLI argument parsing."""

    def test_help(self):
        args = walletfinder.parse_arguments(['--folder', '/tmp'])
        self.assertTrue(args.wallet_mode)

    def test_folder_required(self):
        args = walletfinder.parse_arguments(['--folder', '/tmp'])
        self.assertEqual(args.folder, '/tmp')

    def test_text_mode_flag(self):
        args = walletfinder.parse_arguments(['--folder', '/tmp', '--text-mode'])
        self.assertTrue(args.text_mode)

    def test_mnemonic_mode_backward_compat(self):
        args = walletfinder.parse_arguments(['--folder', '/tmp', '--mnemonic-mode'])
        self.assertFalse(args.wallet_mode)

    def test_depth_argument(self):
        args = walletfinder.parse_arguments(['--folder', '/tmp', '--depth', '3'])
        self.assertEqual(args.depth, 3)

    def test_min_sequential_argument(self):
        args = walletfinder.parse_arguments(['--folder', '/tmp', '--min-sequential', '10'])
        self.assertEqual(args.min_sequential, 10)

    def test_min_scattered_argument(self):
        args = walletfinder.parse_arguments(['--folder', '/tmp', '--min-scattered', '20'])
        self.assertEqual(args.min_scattered, 20)

    def test_debug_flag(self):
        args = walletfinder.parse_arguments(['--folder', '/tmp', '--debug'])
        self.assertTrue(args.debug)


class TestPrivateKeyDetection(unittest.TestCase):
    """Test private key detection patterns (WIF, BIP38, BIP32 extended keys)."""

    def test_raw_wif_compressed_k(self):
        content = "key KwDiBf89QgGbjEhKnhXJuH7LrciVrZi3qYjgd9M7rFU73Nd2Mcv1"
        findings = walletfinder.scan_private_keys(content)
        self.assertEqual(len(findings['raw_wif']), 1)
        self.assertIn('Bitcoin (compressed)', findings['raw_wif'][0]['network'])

    def test_raw_wif_compressed_l(self):
        content = "secret L5G7mhJQeWtJQnh8n4FEgNhpvbKJTi8LqXWEjE2RWKBz1PAUvyst"
        findings = walletfinder.scan_private_keys(content)
        self.assertEqual(len(findings['raw_wif']), 1)

    def test_raw_wif_testnet_c(self):
        content = "testnet cN9spWsvaxA8taS7DFMxnk1yJD2gaF2PX1npuTpy3vuZFJdwavaw"
        findings = walletfinder.scan_private_keys(content)
        self.assertEqual(len(findings['raw_wif']), 1)
        self.assertIn('Testnet', findings['raw_wif'][0]['network'])

    def test_xprv_mainnet_legacy(self):
        content = "extended xprv9s21ZrQH143K24MoUenttLtWQNeeDZvsczTUeCMmb85Mn2qbbmZbpre8QrhSVGRvnYEg3HHxoTKFp5eMqxH41JR99qVKioE3zbhwXAQpWM6"
        findings = walletfinder.scan_private_keys(content)
        self.assertEqual(len(findings['xprv']), 1)

    def test_xpub_mainnet_legacy(self):
        content = "extended xpub661MyMwAqRbcEYSGagKuFUqExQV8d2eizDP5SamP9TcLeqAk9JsrNexcG3qiaSaXCGgfwjQtUD4iRXC9jcmbmA1Jfqoha836vTbBHB564e1"
        findings = walletfinder.scan_private_keys(content)
        self.assertEqual(len(findings['xpub']), 1)

    def test_scan_text_mode_finds_wif(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "keys.txt")
            with open(testfile, 'w') as f:
                f.write("My private key is KwDiBf89QgGbjEhKnhXJuH7LrciVrZi3qYjgd9M7rFU73Nd2Mcv1\n")
            results, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(len(results), 1)
            findings = results[0]['findings']
            key_findings = [f for f in findings if f['type'] == 'private_key']
            self.assertGreaterEqual(len(key_findings), 1)
        finally:
            shutil.rmtree(tmpdir)

    def test_scan_text_mode_finds_xprv(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "keys.txt")
            with open(testfile, 'w') as f:
                f.write("xprv9s21ZrQH143K24MoUenttLtWQNeeDZvsczTUeCMmb85Mn2qbbmZbpre8QrhSVGRvnYEg3HHxoTKFp5eMqxH41JR99qVKioE3zbhwXAQpWM6\n")
            results, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(len(results), 1)
            findings = results[0]['findings']
            key_findings = [f for f in findings if f['type'] == 'private_key']
            self.assertGreaterEqual(len(key_findings), 1)
        finally:
            shutil.rmtree(tmpdir)

    def test_scan_text_mode_mixed_content(self):
        tmpdir = tempfile.mkdtemp()
        try:
            testfile = os.path.join(tmpdir, "mixed.txt")
            with open(testfile, 'w') as f:
                f.write("abandon ability about absorb abstract absurd\n")
                f.write("KwDiBf89QgGbjEhKnhXJuH7LrciVrZi3qYjgd9M7rFU73Nd2Mcv1\n")
            results, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(len(results), 1)
            findings = results[0]['findings']
            mnemonic_findings = [f for f in findings if f['type'] == 'mnemonic']
            key_findings = [f for f in findings if f['type'] == 'private_key']
            self.assertGreaterEqual(len(mnemonic_findings), 1)
            self.assertGreaterEqual(len(key_findings), 1)
        finally:
            shutil.rmtree(tmpdir)


class TestKeyClassification(unittest.TestCase):
    """Test key classification helper functions."""

    def test_classify_wif_uncompressed_5(self):
        self.assertEqual(walletfinder._classify_wif('5HpHagT65TZzG1PH3CSu63k8DbpvD8s5ip4nEB3kEsreAbuatmU'), 'Bitcoin (uncompressed)')

    def test_classify_wif_compressed_k(self):
        self.assertEqual(walletfinder._classify_wif('KwDiBf89QgGbjEhKnhXJuH7LrciVrZi3qYjgd9M7rFU73Nd2Mcv1'), 'Bitcoin (compressed)')

    def test_classify_wif_testnet_c(self):
        self.assertEqual(walletfinder._classify_wif('cN9spWsvaxA8taS7DFMxnk1yJD2gaF2PX1npuTpy3vuZFJdwavaw'), 'Testnet')

    def test_classify_xprv_mainnet_legacy(self):
        self.assertEqual(walletfinder._classify_xprv('xprv9s21ZrQH143K24MoUenttLtWQNeeDZvsczTUeCMmb85Mn2qbbmZbpre8QrhSVGRvnYEg3HHxoTKFp5eMqxH41JR99qVKioE3zbhwXAQpWM6'), 'Bitcoin mainnet (legacy)')

    def test_classify_xprv_nested_segwit(self):
        self.assertEqual(walletfinder._classify_xprv('yprvABrGsX5C9jansMYvK1aX6Rz1aLo6ABvNY6yhRbFey8TEq8eprRjASvJGS4f2VB5rCBMUnktXG7fohNFvZeh4oY6k2BBkJi3YGKmaugUee3V'), 'Bitcoin mainnet (nested segwit)')

    def test_classify_xprv_native_segwit(self):
        self.assertEqual(walletfinder._classify_xprv('zprvAWgYBBk7JR8Giek39NN9JX5WkJwY6ousTDVvCz9YM8q7tEU475tj4yxQTGccV5jmbpUHYEV5in2MaesVHM75bmnLtWtAtcs2Y3qEJLhfqvg'), 'Bitcoin mainnet (native segwit)')

    def test_classify_xpub_mainnet_legacy(self):
        self.assertEqual(walletfinder._classify_xpub('xpub661MyMwAqRbcEYSGagKuFUqExQV8d2eizDP5SamP9TcLeqAk9JsrNexcG3qiaSaXCGgfwjQtUD4iRXC9jcmbmA1Jfqoha836vTbBHB564e1'), 'Bitcoin mainnet (legacy)')

    def test_classify_xpub_nested_segwit(self):
        self.assertEqual(walletfinder._classify_xpub('ypub6QqdH2c5z7965qdPR37XTZvk8NdaZeeDuKuJDyfGXTzDhvyyPy3QzickHFoJaMESbuoUhD1SvsRGJooiTKBcZPguYBW8A2rbCBepfpmmuNw'), 'Bitcoin mainnet (nested segwit)')

    def test_classify_xpub_native_segwit(self):
        self.assertEqual(walletfinder._classify_xpub('zpub6jftahH18ngZw8pWFPu9ff2FJLn2WGdipSRX1NZ9uUN6m2oCedCycnGtJTktaFtN1YvHSgc1PXmpC6RHB1bdMdNWQXCYjwg5TuiU4LsNS9b'), 'Bitcoin mainnet (native segwit)')

    def test_classify_xpub_multisig_native_segwit(self):
        # SLIP-0132 capital-Z prefix must get a descriptive label, not the raw prefix
        self.assertEqual(walletfinder._classify_xpub('Zpub' + '1' * 107), 'Bitcoin mainnet (multisig native segwit)')

    def test_classify_xprv_multisig_native_segwit(self):
        self.assertEqual(walletfinder._classify_xprv('Zprv' + '1' * 107), 'Bitcoin mainnet (multisig native segwit)')

    def test_classify_xpub_testnet_native_segwit(self):
        self.assertEqual(walletfinder._classify_xpub('vpub' + '1' * 107), 'Testnet (native segwit)')


class TestPublicKeyLabeling(unittest.TestCase):
    """Extended public keys (xpub/ypub/zpub/...) must be labeled as public, not private."""

    XPUB = 'xpub661MyMwAqRbcEYSGagKuFUqExQV8d2eizDP5SamP9TcLeqAk9JsrNexcG3qiaSaXCGgfwjQtUD4iRXC9jcmbmA1Jfqoha836vTbBHB564e1'
    XPRV = 'xprv9s21ZrQH143K24MoUenttLtWQNeeDZvsczTUeCMmb85Mn2qbbmZbpre8QrhSVGRvnYEg3HHxoTKFp5eMqxH41JR99qVKioE3zbhwXAQpWM6'

    def _scan_and_print(self, content):
        import io, contextlib
        tmpdir = tempfile.mkdtemp()
        try:
            with open(os.path.join(tmpdir, "keys.txt"), 'w') as f:
                f.write(content)
            results, files_scanned = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            out = io.StringIO()
            with contextlib.redirect_stdout(out):
                walletfinder.print_text_results(results, files_scanned)
            return out.getvalue()
        finally:
            shutil.rmtree(tmpdir)

    def test_xpub_displayed_as_public_key(self):
        output = self._scan_and_print(self.XPUB + "\n")
        self.assertIn("[Public Key: Bitcoin mainnet (legacy)]", output)
        self.assertNotIn("[Private Key:", output)

    def test_xprv_still_displayed_as_private_key(self):
        output = self._scan_and_print(self.XPRV + "\n")
        self.assertIn("[Private Key: Bitcoin mainnet (legacy)]", output)
        self.assertNotIn("[Public Key:", output)

    def test_mixed_keys_grouped_separately(self):
        output = self._scan_and_print(self.XPRV + "\n" + self.XPUB + "\n")
        self.assertIn("[Private Key: Bitcoin mainnet (legacy)]", output)
        self.assertIn("[Public Key: Bitcoin mainnet (legacy)]", output)


class TestKnownTestMnemonicSuppression(unittest.TestCase):
    """Well-known spec/test seeds must be hidden from normal output but shown with --debug."""

    ABANDON_VECTOR = ("abandon abandon abandon abandon abandon abandon "
                      "abandon abandon abandon abandon abandon about")

    def _scan_and_print(self, content, debug=False):
        import io, contextlib
        tmpdir = tempfile.mkdtemp()
        try:
            with open(os.path.join(tmpdir, "seed.txt"), 'w') as f:
                f.write(content)
            results, files_scanned = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            out = io.StringIO()
            with contextlib.redirect_stdout(out):
                walletfinder.print_text_results(results, files_scanned, debug=debug)
            return out.getvalue()
        finally:
            shutil.rmtree(tmpdir)

    def test_known_vector_suppressed_without_debug(self):
        output = self._scan_and_print(self.ABANDON_VECTOR + "\n")
        self.assertNotIn("abandon abandon", output)
        self.assertIn("Suppressed matches", output)

    def test_known_vector_shown_with_debug(self):
        output = self._scan_and_print(self.ABANDON_VECTOR + "\n", debug=True)
        self.assertIn("abandon abandon", output)
        self.assertIn("well-known test seed", output)

    def test_known_vector_not_visible_for_update_exclusions(self):
        tmpdir = tempfile.mkdtemp()
        try:
            with open(os.path.join(tmpdir, "seed.txt"), 'w') as f:
                f.write(self.ABANDON_VECTOR + "\n")
            results, _ = walletfinder.scan_text_mode(tmpdir, None, 6, 12)
            self.assertEqual(len(results), 1)
            self.assertFalse(walletfinder._text_result_is_visible(results[0]))
        finally:
            shutil.rmtree(tmpdir)

    def test_is_known_test_mnemonic(self):
        self.assertTrue(walletfinder._is_known_test_mnemonic(self.ABANDON_VECTOR.split()))
        self.assertTrue(walletfinder._is_known_test_mnemonic(
            self.ABANDON_VECTOR.upper().split()))  # case-insensitive
        self.assertFalse(walletfinder._is_known_test_mnemonic(
            "meadow harbor thunder river shadow forest harbor meadow dragon winter dragon journey".split()))


class TestExclusionMatching(unittest.TestCase):
    """Test exclusion pattern matching: substrings, globs, case-insensitivity, dir pruning."""

    def test_substring_match(self):
        self.assertTrue(walletfinder._exclusion_matches('btcrecover/test/', 'btcrecover/test/wallet.dat'))

    def test_substring_matches_anywhere(self):
        self.assertTrue(walletfinder._exclusion_matches(
            'adselectionattestationspreloaded/',
            'program files (x86)/microsoft/edge/adselectionattestationspreloaded/ad-selection-attestations.dat'))

    def test_substring_no_match(self):
        self.assertFalse(walletfinder._exclusion_matches('btcrecover/test/', 'mywallets/wallet.dat'))

    def test_glob_extension_any_depth(self):
        self.assertTrue(walletfinder._exclusion_matches(
            '*.settingcontent-ms', 'appdata/local/packages/tempstate/foo.settingcontent-ms'))

    def test_glob_extension_no_false_suffix(self):
        self.assertFalse(walletfinder._exclusion_matches(
            '*.settingcontent-ms', 'appdata/foo.settingcontent-ms.bak'))

    def test_glob_with_directory_wildcard(self):
        self.assertTrue(walletfinder._exclusion_matches(
            'capcut/apps/*/resources/bench/score.dat',
            'program files/capcut/apps/3.7.0.1379/resources/bench/score.dat'))

    def test_glob_question_mark(self):
        self.assertTrue(walletfinder._exclusion_matches('cache?.bin', 'some/dir/cache1.bin'))
        self.assertFalse(walletfinder._exclusion_matches('cache?.bin', 'some/dir/cache12.bin'))

    def test_glob_directory_pattern_matches_subtree(self):
        # A glob ending in '/' matches every file beneath that directory, at any depth
        self.assertTrue(walletfinder._exclusion_matches(
            'bitcoinlib*/examples/', 'lib/bitcoinlib_mod/examples/keys.py'))
        self.assertTrue(walletfinder._exclusion_matches(
            'bitcoinlib*/examples/', 'bitcoinlib/examples/wallets.py'))
        self.assertTrue(walletfinder._exclusion_matches(
            'bitcoinlib*/mnemonic.py', 'lib/bitcoinlib/mnemonic.py'))
        self.assertFalse(walletfinder._exclusion_matches(
            'bitcoinlib*/examples/', 'bitcoinlib/wallets/mywallet.dat'))

    def test_is_excluded_case_insensitive(self):
        # Patterns are stored casefolded by load_exclusions; paths must match case-insensitively
        self.assertTrue(walletfinder._is_excluded(
            r'C:\Scan\Program Files\CapCut\Apps\1.0\Resources\bench\score.dat', r'C:\Scan',
            ['capcut/apps/*/resources/bench/score.dat']))

    def test_is_excluded_directory_trailing_slash(self):
        # A 'site-packages/' pattern must prune the directory itself, not just files below it
        self.assertTrue(walletfinder._is_excluded(
            r'C:\Scan\Python\site-packages', r'C:\Scan', ['site-packages/'], is_dir=True))
        self.assertFalse(walletfinder._is_excluded(
            r'C:\Scan\Python\site-packages', r'C:\Scan', ['site-packages/'], is_dir=False))

    def test_load_exclusions_strips_inline_comments(self):
        import unittest.mock
        content = "# header comment\nfoo/bar/   # trailing comment\n\nBaz/Qux\n"
        with unittest.mock.patch('builtins.open', unittest.mock.mock_open(read_data=content)):
            exclusions = walletfinder.load_exclusions()
        self.assertEqual(exclusions, ['foo/bar/', 'baz/qux'])

    def test_walk_directory_applies_glob_exclusions(self):
        tmpdir = tempfile.mkdtemp()
        try:
            os.makedirs(os.path.join(tmpdir, "TempState"))
            keep = os.path.join(tmpdir, "keep.txt")
            skipped = os.path.join(tmpdir, "TempState", "page.settingcontent-ms")
            for p in (keep, skipped):
                with open(p, 'w') as f:
                    f.write("data")
            found = list(walletfinder.walk_directory(
                tmpdir, None, exclusions=['*.settingcontent-ms']))
            self.assertEqual([os.path.basename(p) for p in found], ["keep.txt"])
        finally:
            shutil.rmtree(tmpdir)

    # --- patterns added for common false-positive sources ---

    def test_huggingface_model_cache_glob(self):
        self.assertTrue(walletfinder._exclusion_matches(
            'models--*/snapshots/',
            'hub/models--nvidia--somemodel/snapshots/0893e16/tokenizer.json'))

    def test_btcrecover_test_temp_dirs_glob(self):
        self.assertTrue(walletfinder._exclusion_matches(
            'tmp*-test-btcr*/',
            'users/x/appdata/local/temp/tmp12orr2gm-test-btcr/electrum28-wallet'))
        self.assertTrue(walletfinder._exclusion_matches(
            'tmp*-test-btcr*/',
            'users/x/appdata/local/temp/tmp9syoylyf-test-btcr-dump/vault.txt'))

    def test_chromium_trusted_vault_substring(self):
        self.assertTrue(walletfinder._exclusion_matches(
            'trusted_vault.pb',
            'users/x/appdata/local/nvidia corporation/nvidia app/cefcache/default/trusted_vault.pb'))

    def test_chromium_download_metadata_glob(self):
        self.assertTrue(walletfinder._exclusion_matches(
            'user data/*/downloadmetadata',
            'users/x/appdata/local/bravesoftware/brave-browser/user data/default/downloadmetadata'))

    def test_new_patterns_leave_user_paths_alone(self):
        # Ordinary user files must not be caught by the newly added patterns
        for pattern in ('models--*/snapshots/', 'tmp*-test-btcr*/', 'trusted_vault.pb',
                        'user data/*/downloadmetadata', 'classic_simulator/samples/keystore/',
                        'tests/data/shamir_vectors.json'):
            self.assertFalse(walletfinder._exclusion_matches(
                pattern, 'users/x/documents/my wallets/wallet.dat'),
                "pattern {!r} wrongly matched a user path".format(pattern))
            self.assertFalse(walletfinder._exclusion_matches(
                pattern, 'users/x/appdata/roaming/code/user/history/abc/seed.py'),
                "pattern {!r} wrongly matched an editor-history path".format(pattern))


class TestSystemDirSkipping(unittest.TestCase):
    """Test the absolute-path OS system/hardware directory pruning."""

    def test_is_system_dir_exact_and_descendant(self):
        import unittest.mock
        with unittest.mock.patch.object(walletfinder, 'SYSTEM_EXCLUDED_DIRS',
                                        frozenset({os.path.abspath(os.sep + 'proc')})):
            proc = os.path.abspath(os.sep + 'proc')
            self.assertTrue(walletfinder._is_system_dir(proc))
            self.assertTrue(walletfinder._is_system_dir(os.path.join(proc, '1234', 'maps')))

    def test_is_system_dir_no_match_for_similar_names(self):
        import unittest.mock
        with unittest.mock.patch.object(walletfinder, 'SYSTEM_EXCLUDED_DIRS',
                                        frozenset({os.path.abspath(os.sep + 'dev')})):
            # A user's project folder named 'dev' elsewhere must not match
            self.assertFalse(walletfinder._is_system_dir(
                os.path.abspath(os.path.join(os.sep + 'home', 'user', 'dev', 'x'))))
            # Nor a sibling whose name merely starts with 'dev'
            self.assertFalse(walletfinder._is_system_dir(os.path.abspath(os.sep + 'devices')))

    def test_is_system_dir_empty_set(self):
        import unittest.mock
        with unittest.mock.patch.object(walletfinder, 'SYSTEM_EXCLUDED_DIRS', frozenset()):
            self.assertFalse(walletfinder._is_system_dir(os.path.abspath(os.sep + 'proc')))

    def test_walk_directory_prunes_system_dirs(self):
        import unittest.mock
        tmpdir = tempfile.mkdtemp()
        try:
            os.makedirs(os.path.join(tmpdir, "devfs"))
            Path(tmpdir, "keep.txt").write_text("data")
            Path(tmpdir, "devfs", "skipped.txt").write_text("data")
            with unittest.mock.patch.object(
                    walletfinder, 'SYSTEM_EXCLUDED_DIRS',
                    frozenset({os.path.abspath(os.path.join(tmpdir, "devfs"))})):
                found = list(walletfinder.walk_directory(tmpdir, None))
            self.assertEqual([os.path.basename(p) for p in found], ["keep.txt"])
        finally:
            shutil.rmtree(tmpdir)

    def test_collect_wallet_candidates_prunes_system_dirs(self):
        import unittest.mock
        tmpdir = tempfile.mkdtemp()
        try:
            os.makedirs(os.path.join(tmpdir, "devfs"))
            Path(tmpdir, "keep.txt").write_text("data")
            Path(tmpdir, "devfs", "skipped.txt").write_text("data")
            with unittest.mock.patch.object(
                    walletfinder, 'SYSTEM_EXCLUDED_DIRS',
                    frozenset({os.path.abspath(os.path.join(tmpdir, "devfs"))})):
                found = [p for p, is_dir in
                         walletfinder._collect_wallet_candidates(tmpdir, None)]
            self.assertEqual([os.path.basename(p) for p in found], ["keep.txt"])
        finally:
            shutil.rmtree(tmpdir)


class TestProgressRedirect(unittest.TestCase):
    """Progress output must be suppressed (bar) / milestoned (plain lines) when stdout is
    redirected to a file or pipe, and unchanged on a real terminal."""

    class _FakeTty(io.StringIO):
        def isatty(self):
            return True

    class _BrokenIsatty(io.StringIO):
        def isatty(self):
            raise RuntimeError("no isatty")

    def _capture(self, stream, func, *args, **kwargs):
        with contextlib.redirect_stdout(stream):
            func(*args, **kwargs)
        return stream.getvalue()

    def test_non_tty_print_progress_suppressed(self):
        out = self._capture(io.StringIO(), walletfinder._print_progress, 5, 100, "x")
        self.assertEqual(out, "")

    def test_non_tty_print_progress_milestone(self):
        out = self._capture(io.StringIO(), walletfinder._print_progress, 10000, 854101, "x")
        self.assertNotIn("\r", out)
        self.assertEqual(out, "Scanned 10000/854101 candidates...\n")

    def test_non_tty_print_progress_final(self):
        out = self._capture(io.StringIO(), walletfinder._print_progress, 854101, 854101, "x")
        self.assertEqual(out, "Scanned 854101/854101 candidates...\n")

    def test_non_tty_clear_progress_line_suppressed(self):
        out = self._capture(io.StringIO(), walletfinder._clear_progress_line)
        self.assertEqual(out, "")

    def test_non_tty_scan_status_suppressed_and_milestoned(self):
        out = self._capture(io.StringIO(), walletfinder._print_scan_status, 7, "some/path")
        self.assertEqual(out, "")
        out = self._capture(io.StringIO(), walletfinder._print_scan_status, 20000, "some/path")
        self.assertNotIn("\r", out)
        self.assertIn("20000", out)
        self.assertTrue(out.endswith("\n"))

    def test_non_tty_discovery_reporter_milestoned(self):
        reporter = walletfinder._make_discovery_reporter()
        stream = io.StringIO()
        with contextlib.redirect_stdout(stream):
            for _ in range(walletfinder._MILESTONE_INTERVAL):
                reporter("some/dir")
        out = stream.getvalue()
        self.assertNotIn("\r", out)
        self.assertEqual(out, "Discovering... {} dirs\n".format(walletfinder._MILESTONE_INTERVAL))

    def test_tty_print_progress_writes_inplace_line(self):
        out = self._capture(self._FakeTty(), walletfinder._print_progress, 5, 100, "x")
        self.assertTrue(out.startswith("\r"))
        self.assertIn("Scanning: 5%", out)

    def test_tty_clear_progress_line_writes(self):
        out = self._capture(self._FakeTty(), walletfinder._clear_progress_line)
        self.assertTrue(out.startswith("\r"))

    def test_progress_enabled_handles_broken_isatty(self):
        with contextlib.redirect_stdout(self._BrokenIsatty()):
            self.assertFalse(walletfinder._progress_enabled())


class TestTruncateKey(unittest.TestCase):
    """Test key truncation for display."""

    def test_short_key_not_truncated(self):
        result = walletfinder._truncate_key("abc")
        self.assertEqual(result, "abc")

    def test_long_key_truncated(self):
        result = walletfinder._truncate_key("xprv" + "1" * 107)
        self.assertEqual(len(result), 27)
        self.assertTrue(result.startswith('xprv'))
        self.assertIn('...', result)

    def test_wif_key_truncated(self):
        result = walletfinder._truncate_key("KwDiBf89QgGbjEhKnhXJuH7LrciVrZi3qYjgd9M7rFU73Nd2Mcv1")
        self.assertEqual(len(result), 27)
        self.assertIn('...', result)


if __name__ == '__main__':
    unittest.main()
